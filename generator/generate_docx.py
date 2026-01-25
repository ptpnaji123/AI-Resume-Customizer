from docx import Document
import os


def generate_resume_docx(resume_text, output_path):
    """
    resume_text: str, AI-generated resume
    output_path: str, path to save .docx file
    """

    doc = Document()

    for line in resume_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    return output_path
